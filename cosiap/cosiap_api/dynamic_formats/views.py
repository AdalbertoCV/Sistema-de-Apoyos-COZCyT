from django.shortcuts import render
from common.views import BasePermissionAPIView
from rest_framework.permissions import IsAuthenticated
from users.permisos import es_admin
from rest_framework.response import Response
from rest_framework import status
from rest_framework.generics import get_object_or_404
from .models import DynamicFormat
from users.models import Solicitante
from .serializer import DynamicFormatSerializer
from django.http import HttpResponse,FileResponse, Http404
from io import BytesIO
from docx import Document
import re
from notificaciones.mensajes import Mensaje
import logging
import os


class FormatoConvenio(BasePermissionAPIView):
    ''' 
    Clase para obtener y actualizar el formato default del convenio
    '''

    serializer_class = DynamicFormatSerializer

    permission_classes_list = [IsAuthenticated, es_admin]
    permission_classes_update = [IsAuthenticated, es_admin]


    def get(self, request, *args, **kwargs):
        ''' Obtenemos la plantilla actual '''

        data = {}
        try:
            plantilla = DynamicFormat.objects.get(nombre="formato_convenio_default")
            serializer = self.serializer_class(plantilla)
            data["formato"] = serializer.data
            return Response(data, status = status.HTTP_200_OK)
        except Exception:
            Mensaje.error(data, 'Formato no encontrado.')
            return Response(data, status = status.HTTP_400_BAD_REQUEST)


    def put(self, request, *args, **kwargs):
        ''' actualizamos la plantilla  '''
        formato = DynamicFormat.objects.filter(nombre="formato_convenio_default").first()
        if formato:
            formato = DynamicFormat.objects.get(nombre="formato_convenio_default")
            serializer = self.serializer_class(formato, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
        else:
            serializer = self.serializer_class(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class FormatoMinuta(BasePermissionAPIView):
    ''' 
    Clase para obtener y actualizar el formato default del convenio
    '''

    serializer_class = DynamicFormatSerializer

    permission_classes_list = [IsAuthenticated, es_admin]
    permission_classes_update = [IsAuthenticated, es_admin]


    def get(self, request, *args, **kwargs):
        ''' Obtenemos la plantilla actual '''

        data = {}
        try:
            plantilla = DynamicFormat.objects.get(nombre="formato_minuta_default")
            serializer = self.serializer_class(plantilla)
            data["formato"] = serializer.data
            return Response(data, status = status.HTTP_200_OK)
        except Exception:
            Mensaje.error(data, 'Formato no encontrado.')
            return Response(data, status = status.HTTP_400_BAD_REQUEST)


    def put(self, request, *args, **kwargs):
        ''' actualizamos la plantilla  '''
        formato = DynamicFormat.objects.filter(nombre="formato_minuta_default").first()
        if formato:
            formato = DynamicFormat.objects.get(nombre="formato_minuta_default")
            serializer = self.serializer_class(formato, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
        else:
            serializer = self.serializer_class(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

            

class FormatoAPIView(BasePermissionAPIView):
    ''' 
    Clase para manejar los formatos dinámicos
    '''

    serializer_class = DynamicFormatSerializer

    # Definición de permisos según la acción
    permission_classes_list = [IsAuthenticated, es_admin]
    permission_classes_create = [IsAuthenticated, es_admin]
    permission_classes_update = [IsAuthenticated, es_admin]
    permission_classes_delete = [IsAuthenticated, es_admin]

    def get(self, request, pk=None, *args, **kwargs):
        formatos = DynamicFormat.objects.all()
        serializer = self.serializer_class(formatos, many=True)
        return Response(serializer.data)

    def post(self, request, *args, **kwargs):
        # Crear un nuevo formato
        serializer = self.serializer_class(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, pk=None, *args, **kwargs):
        # Actualizar un formato existente
        formato = get_object_or_404(DynamicFormat, pk=pk)
        serializer = self.serializer_class(formato, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk=None, *args, **kwargs):
        # Eliminar un formato
        formato = get_object_or_404(DynamicFormat, pk=pk)
        formato.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


# Configuración básica de logging
logging.basicConfig(level=logging.DEBUG)

class DescargarFormatoView(BasePermissionAPIView):
    '''
    Clase para manejar las descargas de los formatos dinámicos en DOCX.
    '''
    permission_classes_list = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        try:
            formato_id = kwargs['pk']
            user = request.user
            solicitante = get_object_or_404(Solicitante, id=user.id)

            # Obtén el formato dinámico
            formato = get_object_or_404(DynamicFormat, id=formato_id)
            
            logging.debug(f'Formato ID: {formato_id}')
            logging.debug(f'Solicitante ID: {solicitante.id}, Nombre: {solicitante.nombre}')

            # Cargar el archivo DOCX del formato
            doc_template_path = formato.template.path
            logging.debug(f'Ruta del archivo DOCX: {doc_template_path}')

            doc = Document(doc_template_path)

            # Reemplazar variables en el texto del DOCX
            for para in doc.paragraphs:
                logging.debug(f'Original paragraph text: {para.text}')
                para.text = self.reemplazar_variables(para.text, solicitante)
                logging.debug(f'Replaced paragraph text: {para.text}')

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        logging.debug(f'Original cell text: {cell.text}')
                        cell.text = self.reemplazar_variables(cell.text, solicitante)
                        logging.debug(f'Replaced cell text: {cell.text}')

            # Guardar el documento modificado en un BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Preparar la respuesta para descargar el archivo DOCX
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="formato_{formato_id}.docx"'
            return response

        except Exception as e:
            logging.error(f"Error occurred: {str(e)}", exc_info=True)
            return HttpResponse(f"Error: {str(e)}", status=500)

    def reemplazar_variables(self, text, solicitante):
        """
        Reemplaza las variables en el texto por los valores del solicitante.
        """
        atributos = {
            '{solicitante.nombre}': solicitante.nombre,
            '{solicitante.ap_paterno}': solicitante.ap_paterno,
            '{solicitante.ap_materno}': solicitante.ap_materno if solicitante.ap_materno else None,
            '{solicitante.curp}': solicitante.curp,
            '{solicitante.RFC}': solicitante.RFC,
            '{solicitante.telefono}': solicitante.telefono,
            '{solicitante.direccion}': solicitante.direccion,
            '{solicitante.codigo_postal}': solicitante.codigo_postal,
            '{solicitante.poblacion}': solicitante.poblacion,
            '{solicitante.municipio}': solicitante.municipio.nombre if solicitante.municipio else None,
            '{solicitante.datos_bancarios.nombre_banco}': solicitante.datos_bancarios.nombre_banco if solicitante.datos_bancarios else None,
            '{solicitante.datos_bancarios.cuenta_bancaria}': solicitante.datos_bancarios.cuenta_bancaria if solicitante.datos_bancarios else None,
            '{solicitante.datos_bancarios.clabe_bancaria}': solicitante.datos_bancarios.clabe_bancaria if solicitante.datos_bancarios else None,
        }

        logging.debug(f'Atributos a reemplazar: {atributos}')

        for key, value in atributos.items():
            if value is None:
                value = ''  
            text = text.replace(key, str(value))


        return text


def descargar_formato(request, pk):
    """
    Vista para descargar un formato según su ID.
    """
    try:
        formato = DynamicFormat.objects.get(id=pk)
        response = HttpResponse(formato.template, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{formato.nombre}"'
        return response 
    except DynamicFormat.DoesNotExist:
        raise Http404("El formato solicitado no existe.")


def descarga_manual(request):
    ''' Descarga del manual de usuario '''
    file_path = os.path.join('static/manual/', 'Manual.pdf')  
    try:
        response = FileResponse(open(file_path, 'rb'), content_type='application/pdf')
        response['Access-Control-Allow-Origin'] = '*'
        return response
    except FileNotFoundError:
        raise Http404("El archivo no existe")